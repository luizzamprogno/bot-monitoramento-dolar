from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service as ChromeService
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.wait import WebDriverWait
from selenium.common.exceptions import *
from selenium.webdriver.support import expected_conditions as EC
from datetime import date
from info import *
from docx import Document as word
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from spire.doc import *
from spire.doc.common import *

def iniciar_driver():
    try:
        chrome_options = Options()

        arguments = [
        '--lang=pt-BR',
        '--window-size=1200,800',
        '--incognito',
        '--disable-infobars'
        '--force-device-scale-factor=0.8'
        ]

        for argument in arguments:
            chrome_options.add_argument(argument)

        chrome_options.add_experimental_option('prefs', {
            'download.prompt_for_download': False,
            'profile.default_content_setting_values.notifications': 2,
            'profile.default_content_setting_values.automatic_downloads': 1,
        })

        driver = webdriver.Chrome(service=ChromeService(ChromeDriverManager().install()), options=chrome_options)
        
        wait = WebDriverWait(
            driver=driver,
            timeout=10,
            poll_frequency=1,
            ignored_exceptions=[
                NoSuchElementException,
                ElementNotVisibleException,
                ElementNotSelectableException
            ]
        )

        return driver, wait

    except WebDriverException as e:
        print(f'Erro ao iniciar o driver: {e}')
        return None, None

def open_url(url):
    try:
        driver, wait = iniciar_driver()
        driver.get(url)

        return driver, wait
    
    except Exception as e:
        print(f'Erro ao abrir a URL: {e}')
        return None, None

def colect_usd(current_usd_xpath, wait):
    try:
        return round(float(wait.until(EC.visibility_of_all_elements_located((By.XPATH, current_usd_xpath)))[0].text.replace(',', '.')),2)

    except TimeoutException as e:
        print('Erro ao obter a cotação do dolar')

def get_current_date(date):
    return date.today().strftime("%d/%m/%y")

def save_screenshot(driver):
    driver.save_screenshot('cotacao.png')

def add_hyperlink(paragraph, url, text):
    # Cria um relacionamento para o hyperlink
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Cria o elemento de hyperlink
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Cria um novo run para o texto do hyperlink
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Define a cor do hyperlink
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    # Define o sublinhado do hyperlink
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    run.append(rPr)
    text_element = OxmlElement('w:t')
    text_element.text = text
    run.append(text_element)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)

    return paragraph

def write_doc_content(current_usd, current_date, url):
    document = word()

    heading = HEADING.format(current_usd, current_date)
    content = CONTENT.format(current_usd, current_date)

    heading_paragraph = document.add_heading(heading, level=0)
    heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragrafo = document.add_paragraph(content)

    add_hyperlink(paragrafo, url, 'Clique aqui para ver a cotação atual')

    document.add_paragraph(PRINT_TEXT)
    document.add_picture('./cotacao.png', width=Inches(6), height=Inches(3.5))
    document.add_paragraph(AUTOR)
    document.save('Cotação atual do dolar.docx')

def convert_pdf():

    pdf_document = Document()
    pdf_document.LoadFromFile('.\Cotação atual do dolar.docx')
    pdf_document.SaveToFile('.\Cotação atual do dolar.pdf', FileFormat.PDF)
    pdf_document.Close()

def main():
    driver, wait = open_url(url)
    current_usd_float = colect_usd(current_usd_xpath, wait)
    current_date = get_current_date(date)
    save_screenshot(driver)
    write_doc_content(current_usd_float, current_date, url)
    convert_pdf()

if __name__ == '__main__':
    main()